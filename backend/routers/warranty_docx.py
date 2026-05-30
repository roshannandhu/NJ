"""
Generates a .docx warranty certificate that exactly matches the physical PDF format
and the All_5_Warranties.docx reference document.
"""
import json
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from fastapi import APIRouter, Body, HTTPException
from fastapi.responses import Response

from database import get_db
from models import WarrantyCertificate

router = APIRouter()

# ── Helpers ──────────────────────────────────────────────────────────────────

FONT_BODY = "Times New Roman"
FONT_HEAD = "Times New Roman"


def _set_cell_border(cell, **kwargs):
    """Set borders on a table cell (top/bottom/left/right)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        val = kwargs.get(edge, "single")
        sz  = kwargs.get(f"{edge}_sz", 4)
        if val == "none":
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "nil")
        else:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _shade_cell(cell, fill_hex="D9D9D9"):
    """Fill a table cell with background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _cell_text(cell, text, bold=False, font=FONT_BODY, size=10, align=WD_ALIGN_PARAGRAPH.LEFT):
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)


def _make_doc():
    """Create a blank document with A4 page size and standard margins."""
    doc = Document()
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)
        section.top_margin    = Cm(1.8)
        section.bottom_margin = Cm(1.8)
    # Remove the empty default paragraph at start
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)
    return doc


def _heading(doc, text, level=0, size=11, bold=True, underline=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=8, space_after=4):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    run = para.add_run(text)
    run.bold = bold
    run.underline = underline
    run.font.name = FONT_HEAD
    run.font.size = Pt(size)
    return para


def _body(doc, text, size=10, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
          space_before=0, space_after=4, indent_cm=0):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    if indent_cm:
        para.paragraph_format.left_indent = Cm(indent_cm)
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    return para


def _bullet(doc, text, size=10, char="–", indent_cm=0.4):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.paragraph_format.space_before   = Pt(0)
    para.paragraph_format.space_after    = Pt(3)
    para.paragraph_format.left_indent    = Cm(indent_cm)
    para.paragraph_format.first_line_indent = Cm(-0.3)
    run = para.add_run(f"{char} {text}")
    run.font.name = FONT_BODY
    run.font.size = Pt(size)
    return para


def _blank_line(doc, size=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(size)
    para.paragraph_format.space_after  = Pt(0)


def _cert_row(doc, label, value="", underline_val=True, size=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(2)
    lbl = para.add_run(f"{label}  ")
    lbl.bold = True
    lbl.font.name = FONT_BODY
    lbl.font.size = Pt(size)
    val = para.add_run(value if value else "")
    val.font.name = FONT_BODY
    val.font.size = Pt(size)
    if underline_val:
        val.underline = True
    return para


def _warranty_period_table(doc, rows=None, n_blank=4):
    """Blank or filled warranty period table (2-col: Series | Duration)."""
    _heading(doc, "Warranty Period", size=10, bold=True, space_before=8, space_after=2)
    tbl = doc.add_table(rows=1 + (len(rows) if rows else n_blank), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr = tbl.rows[0]
    hdr.height = Pt(18)
    _cell_text(hdr.cells[0], "Series",        bold=True, size=10)
    _cell_text(hdr.cells[1], "Duration, Years", bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _shade_cell(hdr.cells[0], "BFBFBF")
    _shade_cell(hdr.cells[1], "BFBFBF")
    data = rows if rows else [("", "") for _ in range(n_blank)]
    for i, (s, d) in enumerate(data):
        row = tbl.rows[i + 1]
        row.height = Pt(18)
        _cell_text(row.cells[0], s, size=10)
        _cell_text(row.cells[1], d, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    # Make columns equal width
    for row in tbl.rows:
        for cell in row.cells:
            cell.width = Cm(8.5)
    return tbl


def _sig_area(doc, has_seal=True):
    """Blank signature + seal area at the bottom."""
    _blank_line(doc, 8)
    para = doc.add_paragraph()
    r1 = para.add_run("Company Name: ")
    r1.bold = True
    r1.font.name = FONT_BODY
    r1.font.size = Pt(10)
    r2 = para.add_run("NJ INDIA Trading Pvt. Ltd.")
    r2.font.name = FONT_BODY
    r2.font.size = Pt(10)
    para2 = doc.add_paragraph()
    r3 = para2.add_run("Address: ")
    r3.bold = True
    r3.font.name = FONT_BODY
    r3.font.size = Pt(10)
    r4 = para2.add_run("Bypass Road, Ramanattukara")
    r4.font.name = FONT_BODY
    r4.font.size = Pt(10)
    para3 = doc.add_paragraph()
    r5 = para3.add_run("Seal & Sign:")
    r5.bold = True
    r5.font.name = FONT_BODY
    r5.font.size = Pt(10)
    _blank_line(doc, 28)


def _cert_details_common(doc, cert_data, customer, template_id="default"):
    _blank_line(doc, 4)
    # Divider line via border on a blank paragraph
    para = doc.add_paragraph()
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)

    address = cert_data.get("siteAddress") or customer.get("address") or ""
    product = cert_data.get("productName") or ""
    color   = cert_data.get("productColor") or ""
    prod_full = f"{product}" + (f" — {color}" if color and color != "N/A" else "")
    date    = cert_data.get("purchaseDate") or ""
    seller  = cert_data.get("sellerName") or ""
    batch   = cert_data.get("batchNo") or ""
    cust    = customer.get("name") or ""

    addr_full = f"{cust}{', ' if cust and address else ''}{address}"

    # 1. Address Row (all templates)
    _cert_row(doc, "Address:", f"  {addr_full}")

    # 2. Product Row (all templates, labels depend on template)
    if template_id in ("heatout", "stone_coated", "ceramic"):
        prod_label = "The name of the sold products (complete, including color)"
    else:
        prod_label = "Product Name & Color"
    _cert_row(doc, f"{prod_label}:", f"  {prod_full}")

    # 3. Batch Number Row (Docke and Ceramic)
    if template_id == "docke":
        _cert_row(doc, "Batch Number:", f"  {batch}")
    elif template_id == "ceramic":
        _cert_row(doc, "Batch Number (see on the packaging):", f"  {batch}")

    # 4. Date Row (all templates EXCEPT Ceramic)
    if template_id != "ceramic":
        _cert_row(doc, "Date:", f"  {date}")

    # 5. Trading Organization (all templates EXCEPT Heatout)
    if template_id != "heatout":
        _cert_row(doc, "Trading Organization:", "  NOUFAL & JABBAR INTERNATIONAL LLP", underline_val=False)

    # 6. Seller's Name & Signature Row (all templates)
    _cert_row(doc, "Seller's Name & Signature", f"  {seller}")

    # 7. Date Row for Ceramic only (renders at bottom)
    if template_id == "ceramic":
        _cert_row(doc, "Date:", f"  {date}")

    _blank_line(doc, 20)


# ─────────────────────────────────────────────────────────────────────────────
#  WARRANTY TYPE GENERATORS
# ─────────────────────────────────────────────────────────────────────────────

def _gen_dynamic_warranty(doc, cert_data, customer, template, cert):
    sections = template.get("sections", [])

    # Create a 2-column borderless table
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Make columns equal width (8.5 cm each)
    for cell in tbl.rows[0].cells:
        cell.width = Cm(8.5)
        _set_cell_border(cell, top="none", bottom="none", left="none", right="none")

    # Left Column: Dear Customer + Left Sections
    cell_l = tbl.rows[0].cells[0]
    p_dear = cell_l.paragraphs[0]
    p_dear.paragraph_format.space_before = Pt(0)
    p_dear.paragraph_format.space_after = Pt(2)
    r_dear = p_dear.add_run("Dear Customer,")
    r_dear.bold = True
    r_dear.font.name = FONT_BODY
    r_dear.font.size = Pt(11)

    opening = template.get("opening", "")
    for line in opening.split('\n'):
        if line.strip():
            p = cell_l.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(line)
            r.font.name = FONT_BODY
            r.font.size = Pt(10)

    # Determine split index
    half = 1 if len(sections) == 3 else (len(sections) + 1) // 2
    left_sections = sections[:half]
    right_sections = sections[half:]

    # Helper to add section to a cell
    def add_section_to_cell(cell, sec, is_first_in_cell):
        if is_first_in_cell:
            p_h = cell.paragraphs[0]
        else:
            p_h = cell.add_paragraph()
            
        p_h.paragraph_format.space_before = Pt(6)
        p_h.paragraph_format.space_after = Pt(4)
        r_h = p_h.add_run(sec.get("title", ""))
        r_h.bold = True
        r_h.font.name = FONT_HEAD
        r_h.font.size = Pt(11)
        
        content = sec.get("content", "")
        for line in content.split('\n'):
            if line.strip():
                if sec.get("isBullets"):
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(3)
                    p.paragraph_format.left_indent = Cm(0.4)
                    p.paragraph_format.first_line_indent = Cm(-0.3)
                    r = p.add_run(f"– {line.replace('•', '').replace('-', '').strip()}")
                    r.font.name = FONT_BODY
                    r.font.size = Pt(10)
                else:
                    p = cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(4)
                    r = p.add_run(line)
                    r.font.name = FONT_BODY
                    r.font.size = Pt(10)

    # Render left sections
    for sec in left_sections:
        add_section_to_cell(cell_l, sec, is_first_in_cell=False)
        
    # Render right sections
    cell_r = tbl.rows[0].cells[1]
    for i, sec in enumerate(right_sections):
        add_section_to_cell(cell_r, sec, is_first_in_cell=(i == 0))

    _blank_line(doc, 4)
                    
    # Heatout table
    if template.get("heatoutTable"):
        _blank_line(doc, 6)
        _body(doc, "Years of use counted from the purchase date", bold=True, size=9, space_after=2)
        _body(doc, "Share of the Warrantor liability in % of the purchase price for the replaced product element and the cost of its installation / cost of repairing the product or an element thereof / price paid for the product which price is to be reimbursed", size=9, align=3, space_after=4)
        tbl = doc.add_table(rows=7, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = 1 # CENTER
        liability = [
            ("0 - 10 years",  "100%"),
            ("10 - 12 years", "50%"),
            ("12 - 18 years", "40%"),
            ("18 - 20 years", "30%"),
            ("20 - 21 years", "20%"),
            ("21 - 25 years", "10%"),
        ]
        _cell_text(tbl.rows[0].cells[0], "Years of Use", bold=True, size=10)
        _cell_text(tbl.rows[0].cells[1], "Warrantor Liability %", bold=True, size=10, align=1)
        _shade_cell(tbl.rows[0].cells[0], "BFBFBF")
        _shade_cell(tbl.rows[0].cells[1], "BFBFBF")
        for i, (yr, pct) in enumerate(liability):
            _cell_text(tbl.rows[i+1].cells[0], yr, size=10)
            _cell_text(tbl.rows[i+1].cells[1], pct, size=10, align=1)

    # Series table
    if template.get("showSeriesTable"):
        _blank_line(doc, 6)
        series_table = template.get("seriesTable", [])
        rows = [(s.get("series", ""), s.get("duration", "")) for s in series_table]
        _warranty_period_table(doc, rows=rows if rows else None, n_blank=4)

    # Common Certificate details
    _blank_line(doc, 8)
    _cert_details_common(doc, cert_data, customer, template_id=template.get("id"))
    
    # Signature block
    _sig_area(doc, has_seal=True)


# ─────────────────────────────────────────────────────────────────────────────
#  API ENDPOINT
# ─────────────────────────────────────────────────────────────────────────────

# Map warranty template id → generator function


@router.post("/api/warranties/{wid}/docx")
def warranty_docx(wid: str, body: dict = Body(...)):
    # Use the data sent from the frontend directly; fall back to DB if body is empty
    if body:
        cert = body
    else:
        db = next(get_db())
        try:
            row = db.query(WarrantyCertificate).filter(WarrantyCertificate.id == wid).first()
            if row is None:
                raise HTTPException(status_code=404, detail="Warranty not found")
            cert = json.loads(row.data)
        finally:
            db.close()

    template    = cert.get("template", {})
    cert_data   = cert.get("certData", {})
    customer    = cert.get("customer", {})
    warranty_no = cert.get("warrantyNo") or cert.get("id") or wid

    doc = _make_doc()

    _gen_dynamic_warranty(doc, cert_data, customer, template, cert)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    cust_name = (customer.get("name") or "Customer").replace(" ", "_")
    filename  = f"NJ_Warranty_{warranty_no}_{cust_name}.docx"

    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
